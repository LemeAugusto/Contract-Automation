from openpyxl import load_workbook
from docx import Document
from datetime import datetime

# Passar as informações da planilha para o arquivo word
planilha_fornecedores = load_workbook('./fornecedores.xlsx')
pagina_fornecedores = planilha_fornecedores['Sheet1'] #type:ignore

# Traçar as informações para todas as linhas da planilha
for linha in pagina_fornecedores.iter_rows(min_row=2, values_only=True):
    fornecedor, endereço, cidade, estado, CEP, telefone, email, setor = linha # Fazendo o Unpacking

    arquivo_word = Document()
    arquivo_word.add_heading('Contrato de Prestação de Serviço', 0)

    texto_contrato = f'''

    Este contrato de prestação de serviços é feito entre {fornecedor}, com endereço em {endereço}, 
    {cidade}, {estado}, CEP {CEP}, doravante denominado FORNECEDOR, e a empresa CONTRATANTE.

    Pelo presente instrumento particular, as partes têm, entre si, justo e acordado o seguinte:

    1. OBJETO DO CONTRATO
    O FORNECEDOR compromete-se a fornecer à CONTRATANTE os serviços/material de acordo com as especificações acordadas, respeitando os padrões de qualidade e os prazos estipulados.

    2. PRAZO
    Este contrato tem prazo de vigência de 12 (doze) meses, iniciando-se na data de sua assinatura, podendo ser renovado conforme acordo entre as partes.

    3. VALOR E FORMA DE PAGAMENTO
    O valor dos serviços prestados será acordado conforme as demandas da CONTRATANTE e a capacidade de entrega do FORNECEDOR. Os pagamentos serão realizados mensalmente, mediante apresentação de nota fiscal.

    4. CONFIDENCIALIDADE
    Todas as informações trocadas entre as partes durante a vigência deste contrato serão tratadas como confidenciais.

    Para firmeza e como prova de assim haverem justo e contratado, as partes assinam o presente contrato em duas vias de igual teor e forma.

    FORNECEDOR: {fornecedor}
    E-mail: {email}

    CONTRATANTE: [NOME CONTRATANTE]
    E-mail: [E-MAIL CONTRATANTE]

    {cidade}, {datetime.now().strftime('%d/%m/%Y')}
    '''

    # Salvar aquele arquivo word em uma pasta específica (contratos)
    arquivo_word.add_paragraph(texto_contrato)
    arquivo_word.save(f'./contratos/contrato_{fornecedor}.docx')

